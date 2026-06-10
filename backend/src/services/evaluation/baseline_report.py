"""
Baseline report generator — creates a structured text/docx report
from one or more eval runs.
"""
from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from ...db.eval_models import EvalDataset, EvalErrorAnalysis, EvalRun

logger = logging.getLogger(__name__)


async def build_report_data(run_ids: list[str], db: AsyncSession) -> dict:
    """Collect all data needed for a report."""
    runs = []
    for rid in run_ids:
        r = await db.get(EvalRun, rid)
        if r:
            runs.append(r)
    if not runs:
        return {}

    # Get dataset info from first run
    first = runs[0]
    ds = await db.get(EvalDataset, first.dataset_id) if first.dataset_id else None

    # Error analysis per run
    error_summaries = {}
    for r in runs:
        rows = (await db.execute(
            select(EvalErrorAnalysis).where(EvalErrorAnalysis.run_id == r.id)
        )).scalars().all()
        by_reason: dict[str, int] = {}
        for a in rows:
            by_reason[a.error_reason] = by_reason.get(a.error_reason, 0) + 1
        error_summaries[r.id] = by_reason

    return {
        "runs": runs, "dataset": ds,
        "error_summaries": error_summaries,
        "generated_at": datetime.utcnow().isoformat(),
    }


def _render_text(data: dict) -> str:
    lines = ["# KG-RAG 评测基线报告", f"生成时间: {data.get('generated_at', '')}", ""]
    ds = data.get("dataset")
    if ds:
        lines += [f"## 数据集：{ds.name} v{ds.version}",
                  f"- 题目总数: {ds.total_count}",
                  f"- 题型分布: {ds.type_distribution}",
                  f"- 难度分布: {ds.difficulty_distribution}", ""]

    lines.append("## 各策略准确率对比")
    lines.append(f"{'策略':<20} {'准确率':>8} {'平均延迟':>10} {'题目数':>8}")
    lines.append("-" * 52)
    for r in data.get("runs", []):
        strategy = (r.config or {}).get("strategy", r.run_name)
        acc = f"{(r.accuracy or 0):.1%}"
        lat = f"{r.avg_latency_ms or 0:.0f}ms"
        lines.append(f"{strategy:<20} {acc:>8} {lat:>10} {r.total_questions:>8}")
    lines.append("")

    lines.append("## 错题归因分析")
    for r in data.get("runs", []):
        strategy = (r.config or {}).get("strategy", r.run_name)
        errors = data["error_summaries"].get(r.id, {})
        if errors:
            lines.append(f"\n### {strategy}")
            for reason, cnt in sorted(errors.items(), key=lambda x: -x[1]):
                lines.append(f"  - {reason}: {cnt} 题")

    lines += ["", "## 优化建议",
              "1. 针对「未检索到相关章节」错误：补充同义词词典，增大 top_k",
              "2. 针对「检索到错误文档」错误：强化 source_doc_id 限定逻辑",
              "3. 针对「LLM编造内容」错误：降低 temperature，增强 Prompt 约束",
              "", "---", "由 KG-RAG 评测系统自动生成"]
    return "\n".join(lines)


def _render_docx(data: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading("KG-RAG 评测基线报告", 0)
        doc.add_paragraph(f"生成时间: {data.get('generated_at', '')}")

        ds = data.get("dataset")
        if ds:
            doc.add_heading(f"数据集：{ds.name} v{ds.version}", 1)
            doc.add_paragraph(f"题目总数: {ds.total_count}")
            doc.add_paragraph(f"题型分布: {ds.type_distribution}")

        doc.add_heading("各策略准确率对比", 1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "策略", "准确率", "平均延迟", "题目数"
        for r in data.get("runs", []):
            row = table.add_row().cells
            row[0].text = (r.config or {}).get("strategy", r.run_name)
            row[1].text = f"{(r.accuracy or 0):.1%}"
            row[2].text = f"{r.avg_latency_ms or 0:.0f} ms"
            row[3].text = str(r.total_questions)

        doc.add_heading("错题归因分析", 1)
        for r in data.get("runs", []):
            strategy = (r.config or {}).get("strategy", r.run_name)
            errors = data["error_summaries"].get(r.id, {})
            if errors:
                doc.add_heading(strategy, 2)
                for reason, cnt in sorted(errors.items(), key=lambda x: -x[1]):
                    doc.add_paragraph(f"{reason}: {cnt} 题", style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return _render_text(data).encode("utf-8")


async def generate_report(run_ids: list[str], fmt: str, db: AsyncSession) -> tuple[bytes, str]:
    """Returns (content_bytes, content_type)."""
    data = await build_report_data(run_ids, db)
    if not data:
        return b"No data", "text/plain"
    if fmt == "docx":
        return _render_docx(data), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return _render_text(data).encode("utf-8"), "text/plain"
