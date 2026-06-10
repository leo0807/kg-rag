"""
规范草稿导出 — DOCX 和 Markdown 格式。
"""
from __future__ import annotations

import io
import re
from typing import IO

_SECTION_ORDER = ["1", "2", "3", "4", "5", "6", "7", "8", "9"]

_SECTION_TITLES = {
    "1": "范围",
    "2": "引用文件",
    "3": "术语和定义",
    "4": "材料",
    "5": "设备",
    "6": "技术要求",
    "7": "工艺规程",
    "8": "检验与试验",
    "9": "标识与记录",
}


def _ordered_sections(sections: dict[str, str]) -> list[tuple[str, str, str]]:
    """返回按章节号排序的 (num, title, content) 列表。"""
    result = []
    for num in _SECTION_ORDER:
        content = sections.get(num, "")
        if content and content.strip():
            title = _SECTION_TITLES.get(num, f"第{num}章")
            result.append((num, title, content))
    # 附加不在标准顺序中的章节
    for num, content in sorted(sections.items()):
        if num not in _SECTION_ORDER and content and content.strip():
            result.append((num, f"第{num}章", content))
    return result


def export_to_markdown(spec_name: str, sections: dict[str, str]) -> str:
    lines = [f"# {spec_name}\n"]
    for num, title, content in _ordered_sections(sections):
        lines.append(f"\n## {num} {title}\n")
        lines.append(content.strip())
        lines.append("")
    return "\n".join(lines)


def export_to_docx(spec_name: str, sections: dict[str, str], output: IO[bytes]) -> None:
    """生成 DOCX，尽量符合 COMAC 规范排版风格。"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 页面设置：A4
        section_obj = doc.sections[0]
        section_obj.page_width  = Cm(21)
        section_obj.page_height = Cm(29.7)
        section_obj.left_margin   = Cm(2.5)
        section_obj.right_margin  = Cm(2.5)
        section_obj.top_margin    = Cm(3)
        section_obj.bottom_margin = Cm(2.5)

        # 标题
        title_para = doc.add_heading(spec_name, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")  # 空行

        for num, title, content in _ordered_sections(sections):
            # 章节标题
            h = doc.add_heading(f"{num}  {title}", level=1)
            h.runs[0].font.size = Pt(12)
            h.runs[0].bold = True

            # 内容段落
            for line in content.strip().split("\n"):
                line = line.strip()
                if not line:
                    continue
                # 子章节标题检测（如 6.1、7.2.3）
                if re.match(r"^\d+\.\d", line):
                    sub_h = doc.add_heading(line, level=2)
                    if sub_h.runs:
                        sub_h.runs[0].font.size = Pt(11)
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.first_line_indent = Cm(0.75)
                    if p.runs:
                        p.runs[0].font.size = Pt(10.5)

            doc.add_paragraph("")  # 章节间空行

        doc.save(output)

    except ImportError:
        # python-docx not installed — fallback to plain text in docx container
        _write_plain_docx_fallback(spec_name, sections, output)


def _write_plain_docx_fallback(spec_name: str, sections: dict[str, str], output: IO[bytes]) -> None:
    """python-docx 不可用时，以 UTF-8 文本写入（供调试）。"""
    text = export_to_markdown(spec_name, sections)
    output.write(text.encode("utf-8"))
