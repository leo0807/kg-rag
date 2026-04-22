"""
DOCX 文档解析器
支持 .docx 格式，按 Word 标题样式或粗体段落划分章节
"""
from __future__ import annotations

import re
import logging
from pathlib import Path
from .parser import (
    _looks_like_toc,
    _trim_front_matter_headings,
    is_likely_section_title,
)

logger = logging.getLogger(__name__)

_DOC_ID_RE  = re.compile(r'(CPS\d+)', re.IGNORECASE)
_VER_RE     = re.compile(r'版本[:：]\s*([A-Z])')
_DATE_RE    = re.compile(r'(\d{4}-\d{2}-\d{2})')

# 只匹配第一行，不受多行内容影响；标题字符集更宽松（允许更多 Unicode）
_NUM_RE = re.compile(
    r'^(\d{1,2}(?:\.\d{1,2}){0,3})[ \t\u3000]+'  # 章节号 + 分隔符（含全角空格）
    r'(.{2,80})'                                   # 标题（宽松：任意字符 2-80 位）
    r'$'
)

def _is_heading(para) -> bool:
    """判断段落是否为标题（Word 样式、手动加粗大字号、或文本以章节号开头）"""
    style = para.style.name if para.style else ""
    if style.startswith("Heading") or style.startswith("标题"):
        return not _looks_like_toc(para.text)
    text = para.text.strip()
    if not text:
        return False
    # 文本第一行以章节号格式开头 → 视为标题行（覆盖自定义样式如 NewNormal2）
    first_line = text.split("\n")[0].strip()
    m = _NUM_RE.match(first_line)
    if m and not _looks_like_toc(text):
        return is_likely_section_title(m.group(1), m.group(2))
    # 加粗大字号
    runs = para.runs
    if not runs:
        return False
    bold_chars = sum(len(r.text) for r in runs if r.bold)
    total_chars = max(len(text), 1)
    if bold_chars / total_chars >= 0.8:
        size = runs[0].font.size
        if size is None or size >= 120000:   # ≥9pt
            return True
    return False


def _section_number(text: str) -> str | None:
    """从标题文本提取章节号（只看第一行，避免多行干扰）"""
    first_line = text.strip().split("\n")[0].strip()
    m = _NUM_RE.match(first_line)
    return m.group(1) if m else None


def _trim_front_matter_sections(sections: list[dict]) -> list[dict]:
    if not sections:
        return sections

    trimmed = _trim_front_matter_headings(sections)
    for idx, section in enumerate(trimmed):
        section["seq_index"] = idx
    return trimmed


def parse_docx(path: Path) -> dict:
    """解析 DOCX 文件，返回与 PDF 解析器兼容的 DocumentSchema dict"""
    from docx import Document as _DocxDoc
    from ..models.schemas import DocumentSchema, SectionSchema

    doc = _DocxDoc(str(path))

    # ── 提取元数据 ────────────────────────────────────────────────────────────
    stem = path.stem
    id_m = _DOC_ID_RE.search(stem)
    doc_id = id_m.group(1).upper() if id_m else stem.upper()[:32]

    # 取前 20 段文本拼合，从中提取版本号和日期
    preamble = "\n".join(p.text for p in doc.paragraphs[:20])
    ver_m  = _VER_RE.search(preamble)
    date_m = _DATE_RE.search(preamble)
    version    = ver_m.group(1)  if ver_m  else ""
    issue_date = date_m.group(1) if date_m else ""

    # 标题：优先文档属性 core_properties，否则取首个非空段落
    title = (doc.core_properties.title or "").strip()
    if not title:
        for p in doc.paragraphs:
            t = p.text.strip()
            if len(t) >= 4 and re.search(r'[\u4e00-\u9fff]', t):
                title = t
                break

    # ── 按标题划分章节 ────────────────────────────────────────────────────────
    sections: list[dict] = []
    current_title: str | None   = None
    current_num:   str          = ""
    current_lines: list[str]    = []
    section_idx:   int          = 0

    def _flush():
        nonlocal section_idx
        if current_title is None:
            return
        content = "\n".join(current_lines).strip()
        if not content:
            return
        # 真实章节号：优先用提取到的编号，fallback 用 "0.{idx}" 避免与数字章节混淆
        num = current_num if current_num else f"0.{section_idx}"
        # chunk_id：点替换为下划线，保证合法标识符
        safe_num = num.replace(".", "_")
        level = len(num.split(".")) if re.match(r'^\d', num) else 1
        sections.append({
            "chunk_id":  f"{doc_id}_{safe_num}",
            "number":    num,
            "title":     current_title,
            "content":   f"{num} {current_title}\n{content}",
            "level":     level,
            "seq_index": section_idx,
        })
        section_idx += 1

    # 目录页特征：文本以 tab+数字 结尾（Word TOC 页码引导）
    _TOC_LINE = re.compile(r'\t\d+\s*$')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 跳过目录条目
        if _TOC_LINE.search(text) or _looks_like_toc(text):
            continue
        if _is_heading(para):
            _flush()
            current_title = text
            current_num   = _section_number(text) or ""
            current_lines = []
        else:
            if current_title is None:
                current_title = "前言"
                current_num   = "0"
            current_lines.append(text)

    _flush()
    sections = _trim_front_matter_sections(sections)

    # 处理表格内容（追加到最近章节）
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text and sections:
            sections[-1]["content"] += "\n" + "\n".join(rows_text)

    if not sections:
        # fallback: 整文档作为单一章节
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        sections.append({
            "chunk_id": f"{doc_id}_1", "number": "1", "title": title or stem,
            "content": full_text, "level": 1, "seq_index": 0,
        })

    if not doc_id:
        raise ValueError(f"无法提取文档编号: {path.name}")

    return DocumentSchema(
        doc_id=doc_id, version=version, title=title,
        issue_date=issue_date, total_sections=len(sections),
        sections=[SectionSchema(**s) for s in sections],
        refs=[],
    )
