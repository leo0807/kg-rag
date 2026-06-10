"""
会话导出服务 — 支持 Markdown / JSON / DOCX 格式
"""
from __future__ import annotations

import json
from datetime import datetime
from io import BytesIO
from typing import Any


def _ts(ts: Any) -> str:
    if isinstance(ts, (int, float)):
        return datetime.fromtimestamp(ts / 1000).strftime("%Y-%m-%d %H:%M")
    if isinstance(ts, str):
        return ts[:16]
    return ""


def _role_label(role: str) -> str:
    return {"user": "用户", "assistant": "AI 助手"}.get(role, role)


def export_json(title: str, messages: list[dict], meta: dict) -> bytes:
    payload = {
        "title": title,
        "exported_at": datetime.now().isoformat(),
        "meta": meta,
        "messages": messages,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def export_markdown(title: str, messages: list[dict], meta: dict) -> bytes:
    lines: list[str] = [
        f"# {title}",
        "",
        f"> 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
    ]
    if meta.get("strategy"):
        lines.append(f"> 检索策略：{meta['strategy']}")
    lines.append("")

    for msg in messages:
        role = msg.get("role", "")
        content = msg.get("content", "")
        ts = _ts(msg.get("timestamp", ""))
        lines.append(f"## {_role_label(role)}" + (f"  _{ts}_" if ts else ""))
        lines.append("")
        lines.append(content)
        # sources
        sources: list[dict] = msg.get("sources", [])
        if sources:
            lines.append("")
            lines.append("**来源：**")
            for s in sources[:5]:
                ref = s.get("doc_id", "") or s.get("section_id", "")
                lines.append(f"- {ref}")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


def export_docx(title: str, messages: list[dict], meta: dict) -> bytes:
    """生成 DOCX；若 python-docx 未安装则降级到 UTF-8 文本流。"""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        if meta.get("strategy"):
            doc.add_paragraph(f"检索策略：{meta['strategy']}")
        doc.add_paragraph("")

        for msg in messages:
            role = msg.get("role", "")
            content = msg.get("content", "")
            ts = _ts(msg.get("timestamp", ""))
            heading = doc.add_heading(f"{_role_label(role)}", level=2)
            if ts:
                run = heading.add_run(f"  {ts}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            doc.add_paragraph(content)
            sources: list[dict] = msg.get("sources", [])
            if sources:
                p = doc.add_paragraph("来源：")
                p.runs[0].bold = True
                for s in sources[:5]:
                    ref = s.get("doc_id", "") or s.get("section_id", "")
                    doc.add_paragraph(ref, style="List Bullet")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # fallback: plain UTF-8 wrapped as bytes
        md = export_markdown(title, messages, meta)
        return md
