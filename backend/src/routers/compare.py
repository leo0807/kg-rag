"""工艺文档对比 — 差异摘要 + 图谱变更 + Word 导出"""
from __future__ import annotations
import difflib, io
from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from neo4j import Driver
from ..core.database import get_driver

router = APIRouter(prefix="/api/compare", tags=["compare"])


class CompareRequest(BaseModel):
    doc_id_a: str
    doc_id_b: str


def _sort_key(num: str) -> list:
    try:
        return [int(p) for p in (num or "").split(".") if p.isdigit()]
    except Exception:
        return [0]


def _fetch_sections(driver: Driver, doc_id: str) -> list[dict]:
    with driver.session() as s:
        rows = s.run("""
            MATCH (d:Document {name: $doc_id})-[:HAS_SECTION]->(sec:Section)
            RETURN COALESCE(sec.number, sec.section_number, '') AS number,
                   sec.title AS title, sec.content AS content
        """, doc_id=doc_id)
        return sorted([dict(r) for r in rows], key=lambda r: _sort_key(r.get("number", "")))


def _entity_nodes(driver: Driver, doc_id: str) -> dict[str, dict]:
    """Return entity nodes (Tool/Material/Process/Constraint) linked to doc sections."""
    with driver.session() as s:
        rows = s.run("""
            MATCH (d:Document {name: $doc_id})-[:HAS_SECTION]->(sec:Section)
                  <-[:MENTIONED_IN]-(e)
            WHERE (e:Tool OR e:Material OR e:Process OR e:Constraint)
            RETURN elementId(e) AS id, e.name AS label, labels(e)[0] AS type
        """, doc_id=doc_id)
        return {
            r["id"]: {"id": r["id"], "label": r["label"] or "", "type": r["type"]}
            for r in rows
        }


def _entity_edges(driver: Driver, node_ids: list[str]) -> list[dict]:
    if not node_ids:
        return []
    with driver.session() as s:
        rows = s.run("""
            MATCH (a)-[r]->(b)
            WHERE elementId(a) IN $ids AND elementId(b) IN $ids
            RETURN elementId(a) AS src, elementId(b) AS dst, type(r) AS rel_type
        """, ids=node_ids)
        return [{"from": r["src"], "to": r["dst"], "type": r["rel_type"]} for r in rows]


def _graph_changes(driver: Driver, doc_id_a: str, doc_id_b: str) -> dict:
    nodes_a = _entity_nodes(driver, doc_id_a)
    nodes_b = _entity_nodes(driver, doc_id_b)
    ids_a, ids_b = set(nodes_a), set(nodes_b)

    added_ids   = ids_b - ids_a
    removed_ids = ids_a - ids_b

    edges_a = _entity_edges(driver, list(ids_a))
    edges_b = _entity_edges(driver, list(ids_b))

    def _ekey(e: dict) -> tuple: return (e["from"], e["to"], e["type"])
    keys_a = {_ekey(e) for e in edges_a}
    keys_b = {_ekey(e) for e in edges_b}

    nodes = []
    for nid, node in nodes_a.items():
        nodes.append({**node, "status": "removed" if nid in removed_ids else "unchanged"})
    for nid, node in nodes_b.items():
        if nid in added_ids:
            nodes.append({**node, "status": "added"})

    edges = (
        [{**e, "status": "added"}     for e in edges_b if _ekey(e) not in keys_a] +
        [{**e, "status": "removed"}   for e in edges_a if _ekey(e) not in keys_b] +
        [{**e, "status": "unchanged"} for e in edges_b if _ekey(e) in keys_a]
    )

    return {
        "nodes": nodes,
        "edges": edges,
        "stats": {
            "added_nodes":     len(added_ids),
            "removed_nodes":   len(removed_ids),
            "unchanged_nodes": len(ids_a & ids_b),
            "added_edges":     sum(1 for e in edges if e["status"] == "added"),
            "removed_edges":   sum(1 for e in edges if e["status"] == "removed"),
        },
    }


@router.post("/documents")
async def compare_documents(req: CompareRequest, driver: Driver = Depends(get_driver)):
    import asyncio
    secs_a, secs_b, graph = await asyncio.gather(
        asyncio.to_thread(_fetch_sections, driver, req.doc_id_a),
        asyncio.to_thread(_fetch_sections, driver, req.doc_id_b),
        asyncio.to_thread(_graph_changes,  driver, req.doc_id_a, req.doc_id_b),
    )

    map_a = {s["number"]: s for s in secs_a}
    map_b = {s["number"]: s for s in secs_b}
    all_nums = sorted(set(map_a) | set(map_b), key=_sort_key)

    added, removed, modified = [], [], []
    for n in all_nums:
        in_a, in_b = n in map_a, n in map_b
        if in_b and not in_a:
            added.append({"number": n, "title": map_b[n]["title"]})
        elif in_a and not in_b:
            removed.append({"number": n, "title": map_a[n]["title"]})
        elif map_a[n]["content"] != map_b[n]["content"]:
            diff_lines = list(difflib.unified_diff(
                (map_a[n]["content"] or "").splitlines(),
                (map_b[n]["content"] or "").splitlines(),
                lineterm="", n=1,
            ))
            modified.append({"number": n, "title": map_b[n]["title"], "diff": diff_lines[:20]})

    return {
        "added":         added,
        "removed":       removed,
        "modified":      modified,
        "stats":         {"added": len(added), "removed": len(removed), "modified": len(modified)},
        "graph_changes": graph,
    }


@router.post("/export-docx")
async def export_compare_docx(req: CompareRequest, driver: Driver = Depends(get_driver)):
    import asyncio
    from docx import Document as DocxDoc
    from docx.shared import RGBColor

    secs_a = await asyncio.to_thread(_fetch_sections, driver, req.doc_id_a)
    secs_b = await asyncio.to_thread(_fetch_sections, driver, req.doc_id_b)
    map_a = {s["number"]: s for s in secs_a}
    map_b = {s["number"]: s for s in secs_b}
    all_nums = sorted(set(map_a) | set(map_b), key=_sort_key)

    doc = DocxDoc()
    doc.add_heading(f"对比报告：{req.doc_id_a}  vs  {req.doc_id_b}", 0)

    for num in all_nums:
        sa, sb = map_a.get(num), map_b.get(num)
        if sa and sb and sa["content"] == sb["content"]:
            continue
        label = "[新增]" if not sa else "[删除]" if not sb else "[修改]"
        title = (sb or sa or {}).get("title", "")
        doc.add_heading(f"§{num} {title} {label}", level=2)
        if sa:
            p = doc.add_paragraph()
            run = p.add_run(f"[{req.doc_id_a}]  {(sa['content'] or '')[:600]}")
            run.font.color.rgb = RGBColor(0xCC, 0x44, 0x44)
        if sb and (not sa or sa["content"] != sb["content"]):
            p = doc.add_paragraph()
            run = p.add_run(f"[{req.doc_id_b}]  {(sb['content'] or '')[:600]}")
            run.font.color.rgb = RGBColor(0x22, 0x99, 0x55)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"compare_{req.doc_id_a}_vs_{req.doc_id_b}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )
