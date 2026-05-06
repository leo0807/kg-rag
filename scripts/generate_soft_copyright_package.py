#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE_DIR = ROOT / "soft_copyright_packages" / "_软著母版模板"


def load_spec(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as fh:
        return json.load(fh)


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p


def set_run_font(run, size_pt: int, font_name: str = "仿宋", bold: bool = True) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def restyle_doc(doc: Document, size_pt: int = 12, font_name: str = "仿宋", bold: bool = True) -> None:
    for para in iter_paragraphs(doc):
        for run in para.runs:
            set_run_font(run, size_pt, font_name=font_name, bold=bold)


def replace_text_everywhere(doc: Document, replacements: dict[str, str]) -> None:
    if not replacements:
        return
    for para in iter_paragraphs(doc):
        text = para.text
        new_text = text
        for old, new in replacements.items():
            if old and old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != text:
            para.text = new_text


def copy_template(template_dir: Path, filename: str, target: Path) -> Document:
    src = template_dir / filename
    shutil.copyfile(src, target)
    return Document(target)


def fill_usage_report(doc: Document, spec: dict) -> None:
    software_name = spec["software_name"]
    version = spec.get("version", "V1.0")
    title = spec.get("software_title", f"{software_name} {version}").strip()
    department = spec.get("software_department", "人工智能工程中心")
    function_description = spec["function_description"]
    project_source = spec["project_source"]
    user_evaluation = spec["user_evaluation"]

    # The master report already has the correct table structure and typography.
    # We only replace the content text so the layout stays fixed.
    replacements = {
        "一种基于GraphRAG架构的CPS知识库智能问答与检索系统 V1.0": title,
        "一种基于GraphRAG架构的CPS知识库智能问答与检索系统": software_name,
        "人工智能工程中心": department,
        "该项目面向航空工艺知识问答与检索场景，集成了前后端最新技术栈（如 TypeScript、React/Next.js、Node.js 与 tRPC 等），能够快速构建一个端到端的知识检索与问答系统原型。项目实现了问题输入、检索召回、证据聚合、答案生成和来源追溯等核心功能，同时支持实时结果展示与多轮交互，帮助用户直观理解检索链路的运作过程。该项目内置丰富的开发与测试工具，如代码规范检查、单元测试及容器部署配置，旨在提升代码质量和开发效率，同时为教学演示和原型验证提供完善支持。项目来源：该项目针对航空工艺知识数字化管理与智能检索需求而设计，通过模拟真实问答、知识追溯与结果验证场景，为知识复用和辅助决策提供新的实现思路。": (
            f"{function_description}项目来源：{project_source}"
        ),
        "该项目面向航空工艺知识问答与检索场景，集成了前后端最新技术栈（如 TypeScript、React/Next.js、Node.js 与 tRPC 等），能够快速构建一个端到端的知识检索与问答系统原型。": function_description,
        "项目来源：该项目针对航空工艺知识数字化管理与智能检索需求而设计，通过模拟真实问答、知识追溯与结果验证场景，为知识复用和辅助决策提供新的实现思路。": f"项目来源：{project_source}",
        "该项目功能完善，成功实现了预定的知识问答与检索目标。它将前沿技术与实际业务流程相结合，大大简化了需求对接与业务逻辑实现，使得知识检索与验证过程更加顺畅、高效。": user_evaluation,
    }
    replacements.update(spec.get("replacements", {}))
    replace_text_everywhere(doc, replacements)
    restyle_doc(doc, size_pt=12)


def fill_text_doc(doc: Document, spec: dict, replacements: dict[str, str]) -> None:
    if spec.get("software_name"):
        replacements = dict(replacements)
        replacements[spec.get("source_old_name", "")] = spec["software_name"]
    replace_text_everywhere(doc, replacements)
    restyle_doc(doc, size_pt=12)


def generate_package(spec: dict, template_dir: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    usage_doc = copy_template(template_dir, "软件使用情况报告_母版.docx", output_dir / "软件使用情况报告.docx")
    fill_usage_report(usage_doc, spec["usage_report"])
    usage_doc.save(output_dir / "软件使用情况报告.docx")

    source_doc = copy_template(template_dir, "软件源代码_母版.docx", output_dir / "软件源代码.docx")
    fill_text_doc(source_doc, spec["source_code"], spec["source_code"].get("replacements", {}))
    source_doc.save(output_dir / "软件源代码.docx")

    design_doc = copy_template(template_dir, "软件设计说明书_母版.docx", output_dir / "软件设计说明书.docx")
    fill_text_doc(design_doc, spec["design_doc"], spec["design_doc"].get("replacements", {}))
    design_doc.save(output_dir / "软件设计说明书.docx")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a 3-file soft copyright package from docx masters.")
    parser.add_argument("--spec", required=True, help="Path to a JSON spec file")
    parser.add_argument("--template-dir", default=str(DEFAULT_TEMPLATE_DIR), help="Path to the master template directory")
    parser.add_argument("--output-dir", default=None, help="Where to write the generated package")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    spec_path = Path(args.spec).expanduser().resolve()
    spec = load_spec(spec_path)
    template_dir = Path(args.template_dir).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else Path(spec["output_dir"]).expanduser().resolve()

    generate_package(spec, template_dir, output_dir)
    print(f"generated: {output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
